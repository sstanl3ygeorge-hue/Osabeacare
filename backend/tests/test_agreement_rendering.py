"""
Regression tests for agreement_document_service.py

Three invariants are verified without any live DB or Supabase calls:

1. No raw placeholder text survives rendering — generated contracts must
   contain none of the template tokens or legacy (insert...) phrases.

2. Worker/admin signatures land on the *last* page, not a hardcoded index.
   (Regression: previously _apply_contract_signatures always used page 6.)

3. Missing required contract fields are surfaced — incomplete employee data
   must produce a logged warning rather than silently rendering a broken
   contract that has "TBC" or unfilled tokens for key fields.
"""
from __future__ import annotations

import io
import logging
import sys
import types
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

# ---------------------------------------------------------------------------
# Path bootstrap so the test can import the backend package without a full
# FastAPI / motor install being required at test-collection time.
# ---------------------------------------------------------------------------
BACKEND_DIR = Path(__file__).resolve().parent.parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))


# ---------------------------------------------------------------------------
# Stub out heavy / network-dependent imports before importing the module
# ---------------------------------------------------------------------------
def _make_stub(name: str):
    mod = types.ModuleType(name)
    sys.modules.setdefault(name, mod)
    return mod


for _stub_name in [
    "motor", "motor.motor_asyncio",
    "supabase_storage",
    "services", "services.pdf_service",
]:
    _make_stub(_stub_name)

# pdf_service stub — get_logo_image returns None (no file needed)
sys.modules["services.pdf_service"].get_logo_image = lambda **_kw: None  # type: ignore

# supabase_storage stubs
_ss = sys.modules["supabase_storage"]
_ss.download_file_from_storage = MagicMock()  # type: ignore
_ss.upload_file_to_storage = MagicMock()      # type: ignore

from agreement_document_service import (  # noqa: E402
    CONTRACT_AGREEMENT_TYPE,
    _assert_no_unresolved_contract_placeholders,
    _apply_contract_signatures,
    _build_contract_replacements,
    _docx_to_blocks,
    _format_job_title,
    _render_pdf,
    _replace_contract_text,
    _resolve_contract_fields,
    _validate_contract_fields,
    build_agreement_rendering,
    CONTRACT_TEMPLATE_DOCX_PATH,
    ContractRenderError,
    REQUIRED_CONTRACT_FIELDS,
    read_employee_agreement_state,
)

# ---------------------------------------------------------------------------
# Shared test fixtures
# ---------------------------------------------------------------------------

_COMPLETE_EMPLOYEE = {
    "id": "emp_test_001",
    "name": "Test Worker",
    "job_title": "Care Worker",
    "contract_start_date": "2025-06-01",
    "continuous_service_date": "2025-06-01",
    "hourly_rate": "12.50",
    "sleep_in_rate": "40.00",
}

_ORG_SETTINGS = {
    "organisation_name": "Osabea Healthcare Solutions Ltd",
    "organisation_address": "19 Station Road, Harlow, CM20 2BB",
}

_ORG_SETTINGS_NO_ADDRESS = {
    "organisation_name": "Osabea Healthcare Solutions Ltd",
}

# Phrases that must NOT appear in a rendered contract.  Covers both the
# canonical {{token}} markers and every legacy (insert ...) variant.
_FORBIDDEN_PHRASES = [
    "{{employee_name}}",
    "{{issue_date}}",
    "{{job_title}}",
    "{{contract_start_date}}",
    "{{continuous_service_date}}",
    "{{hourly_rate}}",
    "{{sleep_in_rate}}",
    "{{company_name}}",
    "{{company_address}}",
    "{{commencement_wording}}",
    "(Insert Employee Name)",
    "(insert name of employee)",
    "(insert date of issue)",
    "(insert job title)",
    "(insert 'will commence' or 'commenced')",
    "(insert date this contract starts)",
    "(insert continuous service date of employment)",
    "(insert amount)",
    "insert amount",
    "June 2023",
    "Page 3 of 5",
    "Company ConfidentialSMT1",
    "iCubeDALPro",
    "iCareServicesGroup",
    "Unit 12, Harrods Road, Harlow, CM19 5BJ",
]


def _render_contract_pdf(employee: dict, org_settings: dict) -> bytes:
    """End-to-end render of the DOCX contract template to PDF bytes."""
    from docx import Document

    fields = _resolve_contract_fields(employee, org_settings)
    doc = Document(str(CONTRACT_TEMPLATE_DOCX_PATH))
    blocks = _docx_to_blocks(doc, lambda text: _replace_contract_text(text, fields))
    return _render_pdf(
        blocks,
        title="Employment Contract",
        subtitle=f"{fields['company_name']} | Test",
        employee_name=fields["full_name"],
        add_generated_header=False,
        add_generated_footer=False,
    )


# ---------------------------------------------------------------------------
# Test 1 — No raw placeholder survives rendering
# ---------------------------------------------------------------------------

class TestNoUnreplacedPlaceholders:
    """Regression: contract PDF must contain no raw token or legacy placeholder."""

    def test_full_employee_no_placeholders(self):
        """A fully-populated employee record must leave no unfilled tokens."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)

        # Extract text from every page using PyPDF2
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)

        for phrase in _FORBIDDEN_PHRASES:
            assert phrase not in all_text, (
                f"Forbidden placeholder found in rendered contract: {phrase!r}"
            )

    def test_employee_name_appears_in_output(self):
        """Sanity: the employee's name must appear somewhere in the rendered PDF."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "Test Worker" in all_text, "Employee name not found in rendered contract"

    def test_company_name_replaced(self):
        """The org's name must appear; the old embedded company name must not."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "iCubeDALPro" not in all_text, "Legacy company name survived rendering"
        assert "Osabea Healthcare Solutions Ltd" in all_text, "Org name missing from contract"

    def test_hourly_rate_replaced(self):
        """Pay rate token must be replaced with the employee's rate value."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "12.50" in all_text, "Hourly rate not found in rendered contract"
        assert "(insert amount)" not in all_text, "Legacy pay-rate placeholder survived"

    def test_contract_contains_expected_live_address_and_rate(self):
        """Regression: address/rate must render exactly from canonical values."""
        employee = {**_COMPLETE_EMPLOYEE, "hourly_rate": "12.98"}
        org = {
            "organisation_name": "Osabea Healthcare Solutions Ltd",
            "company_address": "Suite FA4D\n1 St Faith's Street, Maidstone Kent\nME14 1LH",
        }
        pdf_bytes = _render_contract_pdf(employee, org)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "Suite FA4D" in all_text
        assert "1 St Faith's Street, Maidstone Kent" in all_text
        assert "ME14 1LH" in all_text
        assert "12.98" in all_text

    def test_pdf_template_path_not_used(self):
        """Rendering must use the DOCX template; the PDF-overlay path is dead."""
        import asyncio

        from agreement_document_service import (
            _load_template_bytes,
        )

        mock_db = MagicMock()
        # Return a fresh coroutine on each call so each await gets its own object
        mock_db.contract_templates.find_one = lambda *a, **kw: _coro(None)
        mock_db.org_settings.find_one = lambda *a, **kw: _coro(_ORG_SETTINGS)

        result = asyncio.run(
            _load_template_bytes(mock_db, "contract_acceptance")
        )
        _bytes, _name, path = result
        assert path.suffix.lower() != ".pdf", (
            "Contract template loader returned a PDF path — DOCX must be used."
        )
        assert "zero_hour_contract_template_canonical.docx" in str(path)

    def test_canonical_docx_is_tokenised_source(self):
        """Canonical DOCX should contain new tokens, not raw legacy insert placeholders."""
        from docx import Document
        doc = Document(str(CONTRACT_TEMPLATE_DOCX_PATH))
        parts = [p.text for p in doc.paragraphs]
        for sec in doc.sections:
            parts.extend(p.text for p in sec.header.paragraphs)
            parts.extend(p.text for p in sec.footer.paragraphs)
        text = "\n".join(parts)
        assert "{{employee.full_name}}" in text
        assert "{{contract.issue_date}}" in text
        assert "{{company.legal_name}}" in text
        assert "{{contract.hourly_rate}} per hour" in text
        assert "Logo (if required)" not in text
        assert "iCubeDALPro Limited t/a iCareServicesGroup" not in text
        assert "Unit 12, Harrods Road, Harlow, CM19 5BJ" not in text
        assert "June 2023" not in text
        assert "Page 3 of 5" not in text
        assert "Company ConfidentialSMT1" not in text


def _coro(value):
    """Wrap a value in a coroutine that returns a fixed value when awaited."""
    import asyncio

    async def _inner(*_args, **_kwargs):
        return value


class _FakeCursor:
    def __init__(self, rows):
        self._rows = rows

    async def to_list(self, _length):
        return list(self._rows)


class _FakeCollection:
    def __init__(self, rows=None, one=None):
        self._rows = rows or []
        self._one = one

    def find(self, *_args, **_kwargs):
        return _FakeCursor(self._rows)

    async def find_one(self, *_args, **_kwargs):
        return self._one


@pytest.mark.asyncio
async def test_contract_resolver_uses_generated_contract_when_ack_missing():
    employee_id = "emp_contract_001"
    generated_contract_rows = [
        {
            "id": "gc_1",
            "employee_id": employee_id,
            "status": "pending_signature",
            "template_version": "contract_acceptance_v_1",
            "rendered_contract_pdf_url": "contracts/emp_contract_001/current.pdf",
            "generated_at": "2026-05-01T12:00:00Z",
            "created_at": "2026-05-01T12:00:00Z",
        }
    ]

    class _FakeDb:
        agreement_acknowledgements = _FakeCollection(rows=[])
        generated_contracts = _FakeCollection(rows=generated_contract_rows)
        agreement_submissions = _FakeCollection(rows=[], one=None)
        org_settings = _FakeCollection(rows=[], one={})

    result = await read_employee_agreement_state(
        _FakeDb(),
        {"id": employee_id, "name": "Test Employee"},
        CONTRACT_AGREEMENT_TYPE,
    )

    assert result.get("state_label") != "No contract record found"
    assert result.get("contract_state") in {"pending_signature", "awaiting_worker_signature"}
    assert result.get("rendered_file_url") == "contracts/emp_contract_001/current.pdf"
    assert result.get("file_url") == "contracts/emp_contract_001/current.pdf"

    # Return a fresh coroutine object each time the mock is called.
    # We return the coroutine itself; the caller must await it, which is
    # handled by asyncio.run() below rather than the deprecated get_event_loop.
    return _inner()


# ---------------------------------------------------------------------------
# Test 2 — Signatures land on the last page, not a hardcoded index
# ---------------------------------------------------------------------------

class TestSignaturePagePlacement:
    """Regression: _apply_contract_signatures must overlay on the final page."""

    def _make_multi_page_pdf(self, num_pages: int) -> bytes:
        """Create a minimal multi-page PDF with reportlab."""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas as _canvas

        buf = io.BytesIO()
        c = _canvas.Canvas(buf, pagesize=A4)
        for i in range(num_pages):
            c.drawString(72, 700, f"Page {i + 1} of {num_pages}")
            if i < num_pages - 1:
                c.showPage()
        c.save()
        return buf.getvalue()

    def _signature_overlay_page_count(self, pdf_bytes: bytes) -> int:
        """Return total page count of the PDF after signatures are applied."""
        from PyPDF2 import PdfReader
        signed = _apply_contract_signatures(
            pdf_bytes,
            worker_name="Test Worker",
            worker_signed_at="2025-06-01",
        )
        return len(PdfReader(io.BytesIO(signed)).pages)

    @pytest.mark.parametrize("num_pages", [1, 3, 7, 10, 15])
    def test_signature_preserves_page_count(self, num_pages: int):
        """Signing must never add or remove pages regardless of document length."""
        pdf = self._make_multi_page_pdf(num_pages)
        assert self._signature_overlay_page_count(pdf) == num_pages

    def test_signature_on_last_page_not_page_6(self):
        """For a 3-page doc, signatures must NOT fail or corrupt (old code used idx 6)."""
        pdf = self._make_multi_page_pdf(3)
        # Must not raise even though there is no page index 6
        signed = _apply_contract_signatures(
            pdf,
            worker_name="Jane Doe",
            worker_signed_at="2025-07-01",
        )
        from PyPDF2 import PdfReader
        assert len(PdfReader(io.BytesIO(signed)).pages) == 3

    def test_real_contract_signature(self):
        """Sign the real rendered contract — must succeed and preserve all pages."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader

        original_count = len(PdfReader(io.BytesIO(pdf_bytes)).pages)
        signed = _apply_contract_signatures(
            pdf_bytes,
            worker_name="Test Worker",
            worker_signed_at="2025-06-15",
        )
        signed_count = len(PdfReader(io.BytesIO(signed)).pages)
        assert signed_count == original_count, (
            "Signing changed the page count unexpectedly"
        )


# ---------------------------------------------------------------------------
# Test 3 — Missing required fields are surfaced
# ---------------------------------------------------------------------------

class TestMissingFieldValidation:
    """Regression: broken employee data must raise ContractRenderError."""

    def test_missing_hourly_rate_raises(self):
        """An employee with no hourly rate must raise ContractRenderError."""
        employee = {**_COMPLETE_EMPLOYEE, "hourly_rate": None, "pay_rate": None, "rate": None}
        fields = _resolve_contract_fields(employee, _ORG_SETTINGS)
        with pytest.raises(ContractRenderError, match="hourly_rate"):
            _validate_contract_fields(fields)

    def test_missing_contract_start_date_raises(self):
        """An employee with no start date must raise ContractRenderError."""
        employee = {
            k: v for k, v in _COMPLETE_EMPLOYEE.items()
            if k not in ("contract_start_date", "start_date", "employment_start_date",
                         "job_start_date", "promoted_at", "onboarding_start_date")
        }
        fields = _resolve_contract_fields(employee, _ORG_SETTINGS)
        with pytest.raises(ContractRenderError, match="contract_start_date"):
            _validate_contract_fields(fields)

    def test_all_required_fields_present_no_error(self):
        """A fully populated employee record must not raise."""
        fields = _resolve_contract_fields(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        _validate_contract_fields(fields)  # must not raise

    def test_tbc_placeholder_raises(self):
        """Fields that resolve to 'TBC' must cause ContractRenderError."""
        employee = {
            "id": "emp_test_002",
            "name": "Ghost Worker",
            "job_title": "Support Worker",
            # no date fields, no rates → many TBC values
        }
        fields = _resolve_contract_fields(employee, _ORG_SETTINGS)
        with pytest.raises(ContractRenderError):
            _validate_contract_fields(fields)

    @pytest.mark.parametrize("missing_field", REQUIRED_CONTRACT_FIELDS)
    def test_each_required_field_independently(self, missing_field):
        """Every entry in REQUIRED_CONTRACT_FIELDS must individually raise
        ContractRenderError when that field is absent."""
        fields = _resolve_contract_fields(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        fields[missing_field] = ""  # forcibly blank
        with pytest.raises(ContractRenderError, match=missing_field):
            _validate_contract_fields(fields)


# ---------------------------------------------------------------------------
# Test 4 — Contract quality guards
# ---------------------------------------------------------------------------

class TestContractQualityGuards:
    """Rendered contracts must be free of artifacts, TBC values, and leaked company data."""

    def test_no_logo_artifact_in_rendered_contract(self):
        """'Logo (if required)' must never appear in a rendered contract PDF."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "Logo (if required)" not in all_text, (
            "Logo artifact found in rendered contract PDF"
        )

    def test_no_service_generated_footer_in_rendered_contract(self):
        """Contract PDF must not include service-generated footer text."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "Generated by Osabea Healthcare Solutions worker agreements service." not in all_text
        assert "Generated by worker agreements service." not in all_text

    def test_no_tbc_in_fully_resolved_contract(self):
        """A contract rendered with complete data must contain no 'TBC' strings."""
        pdf_bytes = _render_contract_pdf(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "TBC" not in all_text, (
            "TBC placeholder survived into a fully-resolved contract PDF"
        )

    def test_company_address_is_real_address_not_company_name(self):
        """company_address must be the street address, not the organisation name."""
        fields = _resolve_contract_fields(_COMPLETE_EMPLOYEE, _ORG_SETTINGS)
        addr = fields.get("company_address", "")
        company_name = fields.get("company_name", "")
        assert addr != company_name, (
            f"company_address equals company_name: {addr!r}"
        )
        assert "Osabea Healthcare Solutions Ltd" not in addr, (
            f"Company name leaked into company_address: {addr!r}"
        )
        assert "Station Road" in addr, (
            f"Expected street address in company_address, got: {addr!r}"
        )

    def test_missing_organisation_address_raises_on_render(self):
        """Org settings without an address must block rendering via ContractRenderError."""
        fields = _resolve_contract_fields(_COMPLETE_EMPLOYEE, _ORG_SETTINGS_NO_ADDRESS)
        with pytest.raises(ContractRenderError, match="company_address"):
            _validate_contract_fields(fields)

    def test_org_settings_address_precedence_over_employee_override(self):
        """Org settings address must win over stale employee override values."""
        employee = {
            **_COMPLETE_EMPLOYEE,
            "contract_render_overrides": {"company_address": "Stale Address B"},
        }
        org = {
            "organisation_name": "Osabea Healthcare Solutions Ltd",
            "company_address": "Address A",
        }
        fields = _resolve_contract_fields(employee, org)
        assert fields.get("company_address") == "Address A"

    def test_missing_organisation_name_raises_on_render(self):
        """Org settings without a legal/company name must block rendering."""
        fields = _resolve_contract_fields(_COMPLETE_EMPLOYEE, {"organisation_address": "19 Station Road, Harlow, CM20 2BB"})
        with pytest.raises(ContractRenderError, match="company_name"):
            _validate_contract_fields(fields)

    def test_job_title_is_human_formatted(self):
        assert _format_job_title("healthcare_assistant") == "Healthcare Assistant"
        assert _format_job_title("SENIOR-care_worker") == "Senior Care Worker"

    def test_unresolved_contract_placeholders_block_render(self):
        blocks = [
            {"type": "paragraph", "text": "(insert date of issue)"},
            {"type": "paragraph", "text": "£(insert amount) per hour"},
            {"type": "paragraph", "text": "TBC"},
        ]
        with pytest.raises(ContractRenderError, match="unresolved placeholders"):
            _assert_no_unresolved_contract_placeholders(blocks)


def test_contract_render_uses_shared_branded_header_wrapper():
    """Contract render should invoke the shared logo/header wrapper."""
    import asyncio

    mock_db = MagicMock()
    mock_db.contract_templates.find_one = lambda *a, **kw: _coro(None)
    mock_db.org_settings.find_one = lambda *a, **kw: _coro(_ORG_SETTINGS)

    with patch("agreement_document_service.get_logo_image", return_value=None) as logo_mock:
        rendering = asyncio.run(
            build_agreement_rendering(mock_db, _COMPLETE_EMPLOYEE, CONTRACT_AGREEMENT_TYPE)
        )

    assert rendering.get("pdf_bytes"), "Expected rendered PDF bytes for contract"
    assert logo_mock.called, "Shared logo/header helper was not called for contract render"

    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(rendering["pdf_bytes"]))
    all_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Employment Contract" in all_text
    assert "Prepared for: Test Worker" in all_text


def test_contract_render_does_not_crash_when_logo_missing():
    """Missing logo must not break contract generation."""
    import asyncio

    mock_db = MagicMock()
    mock_db.contract_templates.find_one = lambda *a, **kw: _coro(None)
    mock_db.org_settings.find_one = lambda *a, **kw: _coro(_ORG_SETTINGS)

    with patch("agreement_document_service.get_logo_image", return_value=None):
        rendering = asyncio.run(
            build_agreement_rendering(mock_db, _COMPLETE_EMPLOYEE, CONTRACT_AGREEMENT_TYPE)
        )

    assert isinstance(rendering.get("pdf_bytes"), (bytes, bytearray))
    assert len(rendering["pdf_bytes"]) > 0
